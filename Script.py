import docx
import json
import os
from datetime import datetime

def get_docx_metadata(document):
    """
    Extrait les métadonnées d'un document Word.
    Ne garde que les métadonnées qui ne sont pas vides.

    Args:
        document (docx.document.Document): L'objet document python-docx.

    Returns:
        dict: Un dictionnaire contenant uniquement les métadonnées non vides.
    """
    metadata = {}
    properties = document.core_properties

    # Liste des métadonnées à extraire
    metadata_fields = {
        'title': properties.title,
        'subject': properties.subject,
        'author': properties.author,
        'category': properties.category,
        'keywords': properties.keywords,
        'comments': properties.comments,
        'last_modified_by': properties.last_modified_by,
        'revision': properties.revision,
        'version': properties.version
    }
    
    # Ne garder que les métadonnées non vides
    for key, value in metadata_fields.items():
        if value:  # Si la valeur n'est pas None ou vide
            metadata[key] = value
    
    # Gestion des dates
    if properties.created:
        metadata['created'] = properties.created.isoformat()
    if properties.modified:
        metadata['modified'] = properties.modified.isoformat()
    if properties.last_printed:
        metadata['last_printed'] = properties.last_printed.isoformat()

    return metadata

def extract_text_from_docx(docx_filepath):
    """
    Extrait le texte de tous les paragraphes d'un document Word (.docx).
    Ignore les paragraphes vides.

    Args:
        docx_filepath (str): Le chemin complet du fichier Word à lire.

    Returns:
        tuple: Un tuple contenant (liste de paragraphes, objet document docx).
               Retourne (None, None) si le fichier n'est pas trouvé ou s'il y a une erreur.
    """
    try:
        document = docx.Document(docx_filepath)
        # Ne garder que les paragraphes non vides
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return paragraphs, document
    except FileNotFoundError:
        print(f"Erreur : Le fichier Word '{docx_filepath}' n'a pas été trouvé.")
        return None, None
    except Exception as e:
        print(f"Une erreur est survenue lors de la lecture de '{docx_filepath}' : {e}")
        return None, None

def convert_all_docx_to_single_json(source_directory, output_file):
    """
    Convertit tous les documents Word (.docx) d'un répertoire en un seul fichier JSON.
    Chaque document est stocké avec ses métadonnées et son contenu.

    Args:
        source_directory (str): Le répertoire contenant les fichiers DOCX.
        output_file (str): Le chemin du fichier JSON de sortie.
    """
    # Dictionnaire qui contiendra tous les documents
    all_documents = {}
    
    # Parcourir tous les fichiers .docx du répertoire
    for filename in os.listdir(source_directory):
        if filename.endswith(".docx"):
            docx_filepath = os.path.join(source_directory, filename)
            # Extraire le texte et l'objet document
            extracted_paragraphs, document_obj = extract_text_from_docx(docx_filepath)
            
            if extracted_paragraphs is not None:
                # Extraire les métadonnées (uniquement les non vides)
                metadata = get_docx_metadata(document_obj)
                
                # Ajouter le document au dictionnaire principal
                all_documents[filename] = {
                    "metadata": metadata,
                    "content": {
                        "paragraphs": extracted_paragraphs
                    }
                }
                print(f"Le document '{filename}' a été traité avec succès.")

    try:
        # Sauvegarder tous les documents dans un seul fichier JSON
        with open(output_file, mode='w', encoding='utf-8') as json_file:
            json.dump(all_documents, json_file, indent=4, ensure_ascii=False)
        print(f"\nTous les documents ont été convertis et sauvegardés dans '{output_file}'")
    except Exception as e:
        print(f"Une erreur est survenue lors de la sauvegarde du fichier JSON : {e}")

# --- Programme principal ---
if __name__ == "__main__":
    # Configuration des chemins
    SOURCE_DOCX_DIR = "vos_fichiers_word"  # Dossier contenant les fichiers Word
    OUTPUT_JSON_FILE = "Data.json"  # Fichier JSON de sortie

    # Création d'un fichier DOCX d'exemple si le dossier n'existe pas
    if not os.path.exists(SOURCE_DOCX_DIR):
        os.makedirs(SOURCE_DOCX_DIR)
    try:
        # Création d'un document Word de test
        doc = docx.Document()
        # Définition des métadonnées pour le document de test
        doc.core_properties.title = "Rapport Mensuel"
        doc.core_properties.author = "Jean Dupont"
        doc.core_properties.subject = "Analyse des ventes"
        doc.core_properties.keywords = "ventes, rapport, mensuel, analyse"
        doc.core_properties.comments = "Ce rapport couvre les ventes de Mai 2025."
        doc.core_properties.last_modified_by = "Jane Doe"
        # Ajout du contenu
        doc.add_heading('Rapport d\'Activité Mai 2025', level=1)
        doc.add_paragraph('Ceci est le récapitulatif des performances pour le mois écoulé.')
        doc.add_paragraph('Les ventes ont augmenté de 15% par rapport au mois précédent.')
        # Sauvegarde du document
        doc.save(os.path.join(SOURCE_DOCX_DIR, "rapport_mai.docx"))
        print(f"Fichier DOCX d'exemple créé dans '{SOURCE_DOCX_DIR}' pour le test.")
    except Exception as e:
        print(f"Impossible de créer le fichier DOCX d'exemple. Assurez-vous d'avoir 'python-docx' installé. Erreur: {e}")

    # Conversion de tous les documents en un seul fichier JSON
    convert_all_docx_to_single_json(SOURCE_DOCX_DIR, OUTPUT_JSON_FILE)
    print("\nConversion de tous les documents Word en un seul fichier JSON terminée.")